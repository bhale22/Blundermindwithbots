"""
Blundermind Bot Controls -- Project Brief  (.docx generator)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = 'Blundermind_Bot_Controls_Brief.docx'

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)

# ── Colour palette ────────────────────────────────────────────────────────────
AMBER   = RGBColor(0xEB, 0x8C, 0x00)
BLUE    = RGBColor(0x29, 0x9D, 0xD4)
GREEN   = RGBColor(0x5A, 0xA4, 0x5A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
DGREY   = RGBColor(0x22, 0x22, 0x44)
MGREY   = RGBColor(0x66, 0x66, 0x88)
BLACK   = RGBColor(0x11, 0x11, 0x22)
HEAD_BG = (0x0F, 0x34, 0x60)   # table header bg
ROW1_BG = (0x1A, 0x1A, 0x35)
ROW2_BG = (0x22, 0x22, 0x44)
CODE_BG = (0x0D, 0x11, 0x17)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb_tuple):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_ = '{:02X}{:02X}{:02X}'.format(*rgb_tuple)
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_)
    tcPr.append(shd)

def set_para_bg(para, rgb_tuple):
    """Set shading behind a paragraph (for code blocks)."""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    hex_ = '{:02X}{:02X}{:02X}'.format(*rgb_tuple)
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_)
    pPr.append(shd)

def set_cell_borders(cell, color='444466'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_heading(text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(14)
        run.font.color.rgb = AMBER
    elif level == 2:
        run.font.size  = Pt(11)
        run.font.color.rgb = BLUE
    else:
        run.font.size  = Pt(9.5)
        run.font.color.rgb = GREEN
    return para

def add_body(text, italic=False, colour=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.italic = italic
    run.font.color.rgb = colour or DGREY
    return para

def add_bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Inches(0.25)
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = DGREY
    return para

def add_code(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    set_para_bg(para, CODE_BG)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = GREEN
    return para

def add_note(text):
    return add_body(text, italic=True, colour=MGREY)

def add_hr():
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(2)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '444466')
    pBdr.append(bot)
    pPr.append(pBdr)
    return para

def section_heading(text):
    add_heading(text, level=1)
    add_hr()

def sub(text):
    add_heading(text, level=2)

def sub3(text):
    add_heading(text, level=3)

def make_table(headers, rows, col_widths_in):
    """Build a styled table. headers: list of str. rows: list of list of str."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = Inches(col_widths_in[i])
        set_cell_bg(cell, HEAD_BG)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = AMBER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = ROW1_BG if r_idx % 2 == 0 else ROW2_BG
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            cell = cells[c_idx]
            cell.width = Inches(col_widths_in[c_idx])
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(8)
            run.font.color.rgb = WHITE

    doc.add_paragraph()  # spacing after table
    return table

# =============================================================================
# COVER
# =============================================================================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(30)
cover.paragraph_format.space_after  = Pt(4)
r = cover.add_run('BLUNDERMIND')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = AMBER

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Bot Controls -- Technical Project Brief')
r2.font.size = Pt(13); r2.font.color.rgb = BLUE

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Move Selection Mechanics, Formulas & Architecture')
r3.font.size = Pt(10); r3.font.color.rgb = MGREY; r3.italic = True

add_hr()
add_note('All controls described are available in the bot-control-panel.html sidebar. '
         'Formulas reference src/50-bot-engine.js and src/60-bot-ui.js. '
         'This brief reflects the state of the dev branch as of June 2026.')

doc.add_page_break()

# =============================================================================
# TABLE OF CONTENTS
# =============================================================================
add_heading('Table of Contents', level=1)
toc_items = [
    ('1', 'Architecture Overview -- Move Selection Pipeline'),
    ('2', 'Move Source Layer (Opening Book / Maia3 / Stockfish / LCSF)'),
    ('3', 'Think Time Calculation'),
    ('4', 'Time Pressure Play Degradation (Curves A & B)'),
    ('5', 'Temperature & Move Sampling'),
    ('6', 'Strategic Attractors (Style Gauge)'),
    ('7', 'Piece-Type Attractors'),
    ('8', 'Move Quality Range & Luck'),
    ('9', 'Preset Personalities'),
    ('10', 'Time Controls'),
    ('11', 'Human Behaviour Flags'),
]
make_table(['#', 'Section'], toc_items, [0.4, 5.8])
doc.add_page_break()

# =============================================================================
# SECTION 1
# =============================================================================
section_heading('1. Architecture Overview -- Move Selection Pipeline')
add_body('Every bot move passes through six sequential stages. Each stage can veto, '
         'modify, or short-circuit to skip later stages. The stages run in strict order '
         'so that think time -- computed in Stage 2 -- feeds into every downstream '
         'degradation calculation.')

make_table(
    ['Stage', 'Name', 'What happens'],
    [
        ['1', 'Opening Book',
         'ECO table lookup (preferred) or Lichess Masters API (mainline). If a matching move '
         'is found it is played immediately after a brief book delay (400-1200 ms). '
         'Skipped once the game leaves known lines.'],
        ['2', 'Think Time',
         'Actual think time for this move computed from timing mode + Hustle attractor + '
         'clock pressure. Result fed as thinkSec into all downstream degradation curves.'],
        ['3', 'Engine Query',
         'Rough thinkSec estimate used to pick Maia ELO (curve A). '
         'Maia3 or LCSF queried for candidate move probabilities.'],
        ['4', 'Distribution Cut',
         'Curve B sets upper percentile cutoff of candidate list. '
         'Luck attractor shifts the sampling window up or down.'],
        ['5', 'Attractor Reweighting',
         "Each candidate move's probability is multiplied by exp(logBoost) where logBoost "
         'is the sum of all active strategic attractor signals.'],
        ['6', 'Temperature Sampling',
         'Final temperature applied (base T x phase/pressure modifiers). '
         'One move sampled from reweighted distribution.'],
    ],
    [0.4, 1.2, 4.6]
)
add_body('Key architectural principle: Think time is computed first (Stage 2), and the resulting '
         'thinkSec value drives Stages 3, 4, and 5 via the degradation curves. A Coffeehouse '
         'Hustler thinking 0.3 s on a move sees heavy ELO and distribution degradation on that '
         'specific move -- not an average derived from the remaining clock.')
doc.add_page_break()

# =============================================================================
# SECTION 2
# =============================================================================
section_heading('2. Move Source Layer')
sub('2a. Opening Book')
add_body('Two independent book modes are available. Either can be active simultaneously; '
         'each deactivates permanently the first time it fails to find a candidate move.')
make_table(
    ['Mode', 'Source', 'Behaviour'],
    [
        ['Preferred', 'In-memory ECO table (~3,000 named openings)',
         "Matches game history against ECO lines whose ECO code or family prefix is in the bot's "
         "slot list. Plays the next move from the highest-priority matching line. "
         "Exits the book at maxBookDepth (default 20 half-moves)."],
        ['Mainline', 'Lichess Masters API (live)',
         "Queries masters database for the most popular continuation from the current position. "
         "Falls back to Lichess database if no masters result. Deactivates after first cache miss or API error."],
    ],
    [0.9, 1.8, 3.5]
)
add_note('Book delay: 400-1200 ms random pause to simulate human reading time.')

sub('2b. Maia3 Neural Network')
add_body("Maia3 is a series of neural networks (600-2600 ELO) trained to predict human moves "
         "at each rating level. The bot queries the appropriate ELO model (determined by "
         "degradation curve A) and receives a probability distribution over all legal moves.")
add_bullet('Model selection: effective ELO = pressureEffectiveMaiaEloByThink(thinkSec)')
add_bullet('ELO clamped to [600, 2600]')
add_bullet('Probability distribution filtered by quality range (S.8), then reweighted by attractors (S.6)')

sub('2c. Stockfish (SF) Engine')
add_body("Used when the bot engine mode is set to Stockfish. Skill level 1-20 controlled "
         "by sfEffectiveLevel(). The engine is queried via WebWorker; MultiPV probes "
         "provide complexity scores used by the Chaos attractor.")
add_code(
    "sfEffectiveLevel:\n"
    "  blunderFloor   = max(1, round(startLevel x (1 - blunderLimitCp / 400)))\n"
    "  pressureFloor  = max(1, startLevel - round(timePressureMaxDrop / 50))\n"
    "  floorLevel     = max(blunderFloor, pressureFloor)\n\n"
    "  Weaponizer active AND bot ahead  --> return floorLevel immediately\n"
    "  Curve A present --> spline interpolation in log-time space --> lerp level to floor\n"
    "  Fallback: linear ramp startLevel at 30 s --> floor at 0 s"
)

sub('2d. Level-Controlled Stockfish (LCSF)')
add_body("LCSF uses Stockfish with controlled skill variation rather than Maia probability "
         "distributions. sfPickLevel() distributes calls around the target level:")
add_code(
    "Offset distribution:\n"
    "  -2: var2/2   (e.g. 10% at +-2 => 5% chance)\n"
    "  -1: var1/2   (e.g. 30% at +-1 => 15% chance)\n"
    "   0: 1 - var1 - var2  (remaining probability)\n"
    "  +1: var1/2\n"
    "  +2: var2/2\n"
    "Level clamped to [1, 20]"
)
add_body("var1 (+-1 spread): 0-50%.  var2 (+-2 spread): 0-20%. "
         "Setting both to 0 uses the exact target level every move.")
doc.add_page_break()

# =============================================================================
# SECTION 3
# =============================================================================
section_heading('3. Think Time Calculation')
add_body("Think time is calculated in botThinkTime(moveProbs, clockMs) before the engine "
         "is queried. The result (thinkSec) drives all downstream degradation curves, "
         "so attractors that shorten think time (Hustler) directly increase move-quality "
         "degradation for that specific move.")

sub('3a. Timing Modes')
make_table(
    ['Mode', 'Formula / Behaviour'],
    [
        ['Instant',     'Returns 0 ms. No delay whatsoever.'],
        ['Fixed',       'thinkMs = botFixedDelayMs. Capped at min(fixed, clockMs x 0.08) when '
                        'clock < 30 s. Capped at clockMs - 3 s when flagging is disabled.'],
        ['Pace (default)', 'baseSec = (5 x 60) / pace.  complexity = min(1 + entropy x 0.35, 2.5).\n'
                           'thinkMs = baseSec x complexity x 1000.\n'
                           'Pace slider: 5-120 (moves per 5 min equivalent).'],
        ['Complexity',  'cplx = sfCplxScore if available, else min(1, entropy / 4).\n'
                        'mult = botCplxMin + cplx x (botCplxMax - botCplxMin).\n'
                        'thinkMs = botCplxBase x mult x 1000.\n'
                        'Explicit base / min-mult / max-mult sliders in panel.'],
        ['Mirror',      "avg = rolling average of human move timestamps.\n"
                        "thinkMs = avg x jitter(0.8-1.2) x (1 + mirrorOffsetPct/100).\n"
                        "Falls through to Complexity/Pace when no human moves recorded yet."],
    ],
    [1.2, 5.0]
)

sub('3b. Think Time Modifiers (applied in order)')
make_table(
    ['Modifier', 'Condition', 'Effect'],
    [
        ['Weaponizer shortcut', 'Enabled AND bot clock > opp clock + leadMs',
         'Return 0 ms immediately (skip all delays)'],
        ['Move blink', 'botBehavBlink AND entropy < 0.5 (forced/obvious)',
         '200-500 ms random; skips further calculation'],
        ['Clock-pressure mirror', 'botBehavClockMirror AND oppClock < botClock x 0.6',
         'thinkMs x= 0.5'],
        ['Hustle attractor', 'value v in [-5, +5]',
         'thinkMs x= (1 - v x 0.15).\nAt +5: x0.25 (very fast).  At -5: x1.75 (very slow).'],
        ['Reconsideration', 'botBehavReconsider, 15% random',
         'thinkMs x= 1.5 to 2.5 (hesitation pause)'],
        ['Base jitter', 'Always',
         'thinkMs x= uniform(0.8, 1.2)'],
        ['Clock cap (low clock)', 'clockMs < 30,000 ms',
         'thinkMs = min(thinkMs, clockMs x 0.08)'],
        ['Flagging guard', 'botCanFlag = false',
         'thinkMs = min(thinkMs, max(200, clockMs - 3000))'],
        ['Global cap', 'Always',
         'thinkMs = max(200, min(botThinkCapMs(), thinkMs)).\n'
         'Cap = max(6 s, min(45 s, startClock x 0.02)) -- scales with time control.'],
    ],
    [1.35, 1.75, 3.1]
)
add_note('Position entropy is Shannon entropy over Maia move probabilities: H = -sum(p * log2(p)). '
         'Higher entropy = more complex / branching position.')
doc.add_page_break()

# =============================================================================
# SECTION 4
# =============================================================================
section_heading('4. Time Pressure Play Degradation')
add_body("Two independent user-editable spline curves control how the bot's play quality "
         "degrades as think time per move decreases. Both curves use log-scale interpolation "
         "on the X axis (think time in seconds) so that the visually equal spacing on the "
         "panel matches the mathematical interpolation.")
add_code(
    "Curve interpolation (log-x space):\n"
    "  pts sorted by x. For xSec in [pts[i].x, pts[i+1].x]:\n"
    "  t = (ln(xSec) - ln(pts[i].x)) / (ln(pts[i+1].x) - ln(pts[i].x))\n"
    "  y = pts[i].y + t x (pts[i+1].y - pts[i].y)"
)

sub('Curve A -- ELO Degradation')
add_body("Maps think time (s) --> effective Maia ELO. At full think time the bot queries its "
         "configured ELO tier. As think time shrinks the ELO drops, producing weaker move distributions.")
add_code(
    "effectiveELO = pressureEffectiveMaiaEloByThink(thinkSec)\n"
    "             = evalPressureCurve(curveA, thinkSec)  clamped to [600, 2600]"
)
add_bullet('X axis: think time per move in seconds (log scale, typically 0.1 s - 30 s)')
add_bullet('Y axis: Maia ELO (600 - 2600)')
add_bullet('Default: flat at configured ELO (no degradation) unless curve is set')
add_note('The rough thinkSec estimate (computed before the Maia query) is used for ELO selection. '
         'The precise thinkSec (after engine response) is used for curve B and temperature.')

sub('Curve B -- Distribution Cutoff')
add_body("Maps think time (s) --> upper percentile cutoff of the move candidate list. "
         "100% = full distribution. As think time falls the cutoff drops, forcing the bot "
         "to pick from a narrower top slice (higher-probability, safer-looking moves).")
add_code(
    "effectiveUpperPct = pressureEffectiveDayUpperByThink(thinkSec)\n"
    "                  = min(botDayUpper, evalPressureCurve(curveB, thinkSec))"
)
add_bullet('X axis: think time per move in seconds (log scale)')
add_bullet('Y axis: upper percentile 0-100%')
add_bullet('Combined with the Move Quality Range lower bound and Luck shift (S.8)')

sub('Time Pressure Temperature Boost')
add_body("In addition to ELO and distribution effects, low think time raises the sampling "
         "temperature, making move selection more random under pressure:")
add_code(
    "timePressureTempByThink(baseTemp, thinkSec):\n"
    "  boost = { steady: 0.0, normal: 1.0, panicky: 2.5 }[botTimePressure]\n"
    "  if curveB present:\n"
    "    fraction = 1 - evalPressureCurve(curveB, thinkSec) / 100\n"
    "  else:\n"
    "    fraction = max(0, 1 - thinkSec / 30)   # linear fallback\n"
    "  effectiveTemp = baseTemp + fraction x boost"
)
add_note('botTimePressure modes: "steady" (no boost), "normal" (+1.0 max), "panicky" (+2.5 max).')
doc.add_page_break()

# =============================================================================
# SECTION 5
# =============================================================================
section_heading('5. Temperature & Move Sampling')
add_body("After attractor reweighting, one move is sampled from the distribution using "
         "temperature-scaled probabilities. Temperature T controls how peaked or flat "
         "the distribution is before sampling.")
add_code(
    "sampleFromProbs(moveProbs, T):\n"
    "  scaled[m] = prob[m] ^ (1/T)\n"
    "  total = sum(scaled)\n"
    "  sample r ~ Uniform(0, total)\n"
    "  return first m where cumulative sum >= r"
)
add_body("At T=1.0 the original Maia probabilities are used directly. T < 1 sharpens the "
         "distribution (top move more likely). T > 1 flattens it (unlikely moves become "
         "relatively more likely). T is clamped to >= 0.1 to prevent division issues.")

sub('5a. Base Temperature Presets')
make_table(
    ['Preset', 'T value', 'Behaviour'],
    [
        ['Focused', '0.4', 'Strongly favours top 1-2 Maia moves'],
        ['Neutral',  '1.0', "Samples Maia's full distribution as-is"],
        ['Varied',   '1.8', 'Wider sample -- more surprise choices'],
        ['Wild',     '3.0', 'Even low-probability moves get meaningful weight'],
    ],
    [1.2, 0.8, 4.2]
)

sub('5b. Hustler Phase Temperature')
add_body("When the Coffeehouse Hustler personality is active (personalityId = 'hustler'), "
         "the base temperature is replaced by a phase-sensitive curve:")
add_code(
    "hustlerPhaseTemp():\n"
    "  piecesLeft = count of non-pawn, non-king pieces on board\n"
    "  fraction   = clamp(1 - (piecesLeft - 4) / 10, 0, 1)\n"
    "               # 0 at 14 pieces (opening), 1 at <=4 pieces (endgame)\n"
    "  T = 5.0 + fraction x (0.6 - 5.0)\n"
    "    = 5.0 in opening  -->  0.6 in endgame"
)
add_note("14 non-pawn/non-king pieces = full material. At 4 or fewer such pieces, "
         "the endgame temperature floor of 0.6 is reached.")

sub('5c. Complexity-Adjusted Temperature')
add_body("For Maia3 mode, the base temperature can optionally be scaled by position "
         "complexity (MultiPV evaluation spread):")
add_code(
    "complexityAdjustedTemp(baseT):\n"
    "  if sfCplxScore is null --> return baseT\n"
    "  cplxFactor = 0.8 + sfCplxScore x 0.4   # 0.8 at simple, 1.2 at complex\n"
    "  return baseT x cplxFactor"
)
add_note("sfCplxScore is the normalized std. deviation of the top-N Stockfish evaluations, "
         "ranging 0 (all moves equal) to 1 (widely spread evaluations).")
doc.add_page_break()

# =============================================================================
# SECTION 6
# =============================================================================
section_heading('6. Strategic Attractors (Style Gauge)')
add_body("Strategic attractors reweight each candidate move's probability by multiplying "
         "it by exp(logBoost). The total logBoost for a move is the sum of all active "
         "attractor contributions. Attractors are zero at centre and scale linearly "
         "with slider value (-5 to +5).")

sub('6a. CP Budget and Scale Factor')
add_body("The CP (centipawn) budget (0-300) is a single master gain knob for all attractors. "
         "It controls how strongly attractors push the distribution, regardless of individual "
         "slider positions.")
add_code(
    "totalAbs = sum |v| for all attractor and piece sliders\n"
    "scale    = cpBudget / (totalAbs x 150)   if cpBudget > 0 and totalAbs > 0\n"
    "         = 0                              otherwise\n\n"
    "For each move:  prob_new = prob x exp(logBoost)\n"
    "  where logBoost = sum( attractor_v x scale x signal )\n\n"
    "Example: cpBudget=150, one slider at 1, all others 0:\n"
    "  scale = 1/150, signal = 1  -->  logBoost = 1  -->  exp(1) = 2.7x boost"
)
add_note("Allocating 150 cp of budget on a single maximally-set attractor produces roughly "
         "a 2.7x probability boost on moves that fully satisfy that attractor.")

sub('6b. Attractor Formulas')
make_table(
    ['Attractor', 'Left (-) / Right (+)', 'Signal formula', 'Notes'],
    [
        ['Chaos Agent\n/ Simplifier',
         '- Simplifier: low-variance moves\n+ Chaos Agent: high-variance moves',
         'logBoost += chaosVal x scale x signal\n(signal from MultiPV sigma)',
         'Requires MultiPV SF probe. sigma = std dev of top-N eval scores.'],
        ['Complicate / Simplify winning',
         '- Simplify when winning\n+ Complicate when winning',
         'Active when eval >= +100 cp. Boosts moves that increase or decrease MultiPV variance.',
         'Only fires in clearly winning positions. Combine with Chaos for full effect.'],
        ['Space Cadet\n/ Space Waster',
         '- Space Waster: cede space\n+ Space Cadet: reduce weak sq.',
         'delta = currentWeakSq - simWeakSq\nlogBoost += v x scale x tanh(delta/5)',
         'Full board attack map on simulated position. Captures discovered attacks and blocking moves.'],
        ['Fort Knox\n/ Glass Cannon',
         '- Glass Cannon: exposed pieces\n+ Fort Knox: more defenders',
         'delta = simTotalDefs - currentDefs\nlogBoost += v x scale x tanh(delta/3)',
         'Sums defensive coverage across all bot pieces. Full board attack map.'],
        ['Gambito\n/ Gambit Shy',
         '- Gambit Shy: avoid sacrifices\n+ Gambito: follow gambit lines',
         'Opening only (< 20 half-moves).\nECO match: logBoost += v x scale.\nFallback: pawn to attacked, undefended sq.',
         'ECO gambit line match takes priority. Structural fallback if ECO data not yet loaded.'],
        ['Trade Seeker\n/ Trade Avoider',
         '- Trade Avoider: no captures\n+ Trade Seeker: captures & threats',
         'Capture: logBoost += v x scale\nThreat: logBoost += v x scale\n  x tanh(newThreats / 2)',
         'Non-captures scored by new opponent-piece threats created from destination square.'],
        ['Rigid / Loose\nPawn Structure',
         '- Loose: open mobile structure\n+ Rigid: tighter structure',
         'Own pawn moves only.\ndelta = currentPenalty - simPenalty\nlogBoost += v x scale x tanh(delta)',
         'Penalty = islands + doubled + isolated pawns. delta > 0 = move tightens structure.'],
        ['Coffeehouse Hustler\n/ Overthinker',
         '- Overthinker: longer think time\n+ Hustler: shorter think time',
         'Applied to think time:\nthinkMs x= (1 - v x 0.15)',
         'At +5: x0.25 (4x faster). At -5: x1.75. Feeds into pressure curves as actual thinkSec.'],
        ['Luck / Bad Day',
         '- Good day: top of distribution\n+ Bad day: bottom of distribution',
         'Shifts quality range window:\nlo = botDayLower - v x 4\nhi = pressureUpper - v x 4',
         'See S.8 Move Quality Range for full description.'],
        ['Panicky\n/ Calm under pressure',
         '- Calm: no temperature boost\n+ Panicky: large T boost under pressure',
         'Modifies timePressure boost mode.\n(steady / normal / panicky)',
         'See S.4 Pressure Temperature Boost.'],
        ['Attacker\n/ Peacemaker',
         '- Peacemaker: quiet positions\n+ Attacker: maximize threats',
         'totalThreats = sum attacks on each opp. piece\nlogBoost += v x scale x tanh(threats/6)',
         'Full board attack map on simulated position. Counts attack coverage across all opp. pieces.'],
    ],
    [1.0, 1.4, 1.8, 2.0]
)
add_note("All attractor signals using tanh map their input to a smooth -1..+1 range, "
         "preventing any single move from receiving an unbounded boost.")
doc.add_page_break()

# =============================================================================
# SECTION 7
# =============================================================================
section_heading('7. Piece-Type Attractors')
add_body("Six independent sliders (-5 to +5) bias the bot toward or away from moving "
         "each piece type. The formula is the same for all six:")
add_code("logBoost += pieceVals[pieceType] x scale")
add_body("The piece type is determined from the moving piece before the move. "
         "The signal is 1 (no continuous function as with strategic attractors). "
         "Scale is shared with strategic attractors so the CP budget controls all attractor "
         "strength together.")

make_table(
    ['Piece', 'Left (-) behaviour', 'Right (+) behaviour', 'Best combined with'],
    [
        ['Pawn',   'Suppress pawn moves',              'Prioritise pawn advances',        'Rigid structure (connected chains)'],
        ['Knight', 'De-prioritise knight hops',        'Knight manoeuvres first',          'Rigid structure (closed positions)'],
        ['Bishop', 'Keep bishops passive',             'Diagonal play, bishop pair',       'Loose structure (open diagonals)'],
        ['Rook',   'Passive rooks, avoid early trades','File seizure, rook lifts',         'Open files, trade-seeker'],
        ['Queen',  'Conservative queen',               'Active queen, queen-led attacks',  'Chaos Agent (keeps complications)'],
        ['King',   'King safety, stay castled',        'King activity, endgame aggression','Endgame positions, low piece count'],
    ],
    [0.7, 1.5, 1.5, 2.5]
)
doc.add_page_break()

# =============================================================================
# SECTION 8
# =============================================================================
section_heading('8. Move Quality Range & Luck')
add_body("The Move Quality Range dual slider selects which segment of the Maia "
         "probability-ranked candidate list the bot samples from. Candidates outside "
         "the [lower, upper] percentile window are excluded before temperature sampling "
         "and attractor reweighting.")

sub('8a. Percentile Band Filtering')
add_code(
    "sorted candidates by probability (highest first)\n"
    "cumulative probability sum = total\n\n"
    "band.lo = (lo/100) x total\n"
    "band.hi = (hi/100) x total\n\n"
    "Include move m if:\n"
    "  cumulative_end(m) > band.lo  AND  cumulative_start(m) < band.hi"
)
add_body("Default: lo=0%, hi=100% (full distribution). Setting lo=20%, hi=60% samples "
         "only the middle tier -- avoiding the top moves (too accurate) and the worst moves "
         "(random blunders).")

sub('8b. Luck Attractor Shift')
add_body("The Luck attractor shifts both bounds of the quality window simultaneously:")
add_code(
    "lo_effective = botDayLower - luckVal x 4\n"
    "hi_effective = pressureUpper - luckVal x 4\n\n"
    "luckVal > 0 (Bad day):  window shifts down  --> samples lower-ranked moves  --> worse play\n"
    "luckVal < 0 (Good day): window shifts up    --> samples top-ranked moves    --> sharper play"
)
add_note("pressureUpper is the curve-B degraded upper bound, so luck shift compounds with "
         "time pressure effects.")

sub('8c. Blunder Limit and Min-Probability Filter')
add_body("Two additional filters prune the candidate list before the quality range is applied:")
add_code(
    "relFloor  = bestProb x exp(-blunderLimitCp / 100)\n"
    "absFloor  = minProbPct / 100\n"
    "threshold = max(absFloor, relFloor)\n"
    "keep only moves with prob >= threshold"
)
add_bullet("blunderLimitCp (0-400 cp): sets how far below the best move a candidate can fall. "
           "50 cp = tight;  400 cp = everything allowed.")
add_bullet("minProbPct (0-10%): absolute probability floor; removes near-zero candidates.")
doc.add_page_break()

# =============================================================================
# SECTION 9
# =============================================================================
section_heading('9. Preset Personalities')
add_body("Each preset personality is a named collection of attractor values. Loading a "
         "personality sets all attractor sliders to the preset values. Values use the "
         "same -5 to +5 scale as manual sliders.")

make_table(
    ['Personality', 'Tag', 'Notable attractor values', 'Character'],
    [
        ['Captain Entropy', 'Chaos',
         'chaos:+4, compwin:+4, gambito:+3\ntrade:+2, spacecadet:+2\nhustle:-3, structure:-3',
         'Seeks complications, gambit lines, and tactical chaos. Thinks longer per move (hustle -3).'],
        ['Norm', 'Order',
         'structure:+4, fortkx:+3\nchaos:-4, compwin:-4\ngambito:-2, hustle:+2',
         'Closes positions, defends solidly. Simplifies when winning. Slightly faster pace.'],
        ['Attacky McTackerson', 'Captures',
         'trade:+4, attacker:+3, chaos:+2\nhustle:-2, structure:-1',
         'Favours captures and threats. Slightly deliberate (hustle -2).'],
        ['Overthinker', 'Methodical',
         'structure:+3, spacecadet:+3, fortkx:+2\nhustle:-4, pressure:-3\nchaos:-3, compwin:-2',
         'Maximises board control and structure. Thinks long on every move (hustle -4). '
         'Likely to accumulate time pressure in classical games.'],
        ['Coffeehouse Hustler', 'Fast',
         'hustle:+5, pressure:+2, luck:+2\ngambito:+2, structure:-2',
         'Plays fast and loose. Phase temperature: T=5 opening --> T=0.6 endgame. '
         'At hustle +5: think time x0.25 of baseline.'],
        ['The Blunderer', 'Human',
         'luck:+3, pressure:+5',
         'Plays normally but cracks unpredictably under pressure. '
         'Luck shifts toward lower-ranked moves.'],
        ['The Hoarder', 'Material',
         'gambito:-3, trade:-3, fortkx:+2\nchaos:-2, hustle:+2, luck:-2',
         'Never sacrifices material. Avoids trades and gambits. Slight faster pace.'],
        ['Pawn Chain Gang', 'Pawns',
         'structure:+5, fortkx:+1, hustle:+1\nchaos:-1, compwin:-2',
         'Maximises pawn structure. Slightly faster pace.'],
        ['Spite Check', 'Checks',
         'chaos:+3, compwin:+2, attacker:+3\ntrade:+2, hustle:-2',
         'Prefers checking moves. Builds attacks. Slower pace (hustle -2).'],
        ['Clock Watcher', 'Clock',
         'compwin:-2, trade:-1, pressure:-3',
         'Conservative when ahead on time, reckless when behind. Calm under pressure.'],
        ['Grandmaster Bad Day', 'Human',
         'luck:+5, pressure:+2',
         'Maia 2400 ELO but samples deep in distribution -- strong player having an off day.'],
    ],
    [1.3, 0.65, 2.1, 2.15]
)
add_note("Hustle attractor sign convention: +5 = Coffeehouse Hustler (fast, thinkMs x0.25). "
         "-5 = Overthinker (slow, thinkMs x1.75).")
doc.add_page_break()

# =============================================================================
# SECTION 10
# =============================================================================
section_heading('10. Time Controls')
add_body("Time controls are defined in TIME_CONTROLS and selected from the bot panel grid. "
         "The grid is divided into Blitz (<=5 min), Rapid (10-30 min), "
         "Classical/Correspondence (60 min+), Infinite (untimed), and a special "
         "90+30 column for FIDE/Candidates format.")

make_table(
    ['Format', 'Base time', 'Increment', 'Bonus', 'Bonus at move', 'Inc from move'],
    [
        ['Bullet 1+0',      '1 min',  '0 s',  '--', '--', '--'],
        ['Blitz 3+2',       '3 min',  '2 s',  '--', '--', '--'],
        ['Blitz 5+0',       '5 min',  '0 s',  '--', '--', '--'],
        ['Rapid 10+0',      '10 min', '0 s',  '--', '--', '--'],
        ['Rapid 15+10',     '15 min', '10 s', '--', '--', '--'],
        ['Classical 30+0',  '30 min', '0 s',  '--', '--', '--'],
        ['Classical 60+0',  '60 min', '0 s',  '--', '--', '--'],
        ['Classical 90+0',  '90 min', '0 s',  '--', '--', '--'],
        ['FIDE 90+30 *',    '90 min', '30 s', '+30 min', 'Move 40', 'Move 41'],
        ['Untimed',         'Inf',    '0 s',  '--', '--', '--'],
    ],
    [1.2, 0.8, 0.8, 0.8, 0.95, 0.95]
)
add_note("* 90+30 (FIDE / Candidates): 90 minutes for first 40 moves, then +30 minutes added "
         "to both clocks at move 40. 30-second increment applies from move 41 onward. "
         "The clockBonusApplied flag prevents the bonus from being awarded twice.")

sub('Bonus Time Implementation')
add_code(
    "clockAfterMove():\n"
    "  fullMoveNum = ceil(gameMovesAlgebraic.length / 2)\n"
    "  if tc.bonusSecs AND fullMoveNum >= tc.bonusAtMove AND NOT clockBonusApplied:\n"
    "    clockTimeW += bonusSecs  (capped at 59940 s)\n"
    "    clockTimeB += bonusSecs\n"
    "    clockBonusApplied = true\n"
    "  if clockInc > 0 AND fullMoveNum >= tc.incFromMove:\n"
    "    add increment to side that just moved"
)
doc.add_page_break()

# =============================================================================
# SECTION 11
# =============================================================================
section_heading('11. Human Behaviour Flags')
add_body("Human behaviour flags add realistic timing irregularities and clock-pressure "
         "responses. Each flag is independently toggleable. They modify think time or "
         "move selection in specific situations.")

make_table(
    ['Flag', 'When active', 'Effect'],
    [
        ['Move Blink\n(botBehavBlink)',
         'Maia3 mode only.\nPosition entropy < 0.5\n(forced or obvious move)',
         'Returns 200-500 ms random pause immediately, skipping full think time calculation. '
         'Simulates a human instantly playing the obvious recapture or forced reply.'],
        ['Reconsider\n(botBehavReconsider)',
         '15% random chance\nper move',
         'Multiplies computed think time by 1.5-2.5x. Simulates the human hesitation '
         'of starting to play a move then reconsidering.'],
        ['Clock Mirror\n(botBehavClockMirror)',
         'Opponent clock <\nbot clock x 0.6',
         "Halves think time when the opponent is significantly lower on time. "
         "Simulates a human speeding up to maintain clock advantage."],
        ['Clock Weaponizer\n(botWeaponizerEnabled)',
         'Bot clock ahead by\nmore than leadMs',
         'Returns 0 ms think time AND drops Stockfish to floor level. '
         'Maximises time pressure on opponent by playing instantly when already ahead on clock. '
         'leadMs configured by Weaponizer Lead Time slider.'],
        ['Can Flag\n(botCanFlag)',
         'Toggle: on or off',
         'When OFF: think time capped at max(200 ms, clockMs - 3000 ms). '
         'Prevents bot from flagging itself. When ON: no such safeguard.'],
    ],
    [1.2, 1.5, 3.5]
)

add_hr()
add_note("End of brief. For implementation details see src/50-bot-engine.js (engine), "
         "src/60-bot-ui.js (config application), and bot-control-panel.html (UI and presets).")
add_note("Generated June 2026 -- Blundermind Bot Controls v1.0")

doc.save(OUT)
print('Written: ' + OUT)
